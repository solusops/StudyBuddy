"""File and text ingestion into ChromaDB.

Chunk type is recorded in metadata so the Tutor Agent can preferentially
query question chunks for quiz/flashcard generation.

Files are deduplicated by SHA-256 content hash — ingest_file() is a no-op if
the file is already indexed, making repeated library scans cheap.
"""
import io
import uuid
from typing import Literal, Optional

from app.rag.chunker import chunk_text
from app.rag.chromadb_client import ChromaDBClient

ChunkType = Literal["content", "question"]

# Single persistent collection name for all library content
LIBRARY_COLLECTION = "library"


def _extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_bytes: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def extract_document_structure(file_bytes: bytes, filename: str, max_chars: int = 6000) -> str:
    """Extract only the structural overview of a document for instant tree generation.

    Returns headings/TOC and first-page text — enough for Gemma 4 to derive a
    curriculum tree without needing full chunking.
    """
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        # First 4 pages capture TOC + intro for most textbooks
        pages = [reader.pages[i].extract_text() or "" for i in range(min(4, len(reader.pages)))]
        return "\n".join(pages)[:max_chars]
    elif ext == "docx":
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        # Only heading paragraphs for structure
        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        if headings:
            return "\n".join(headings)[:max_chars]
        # Fall back to first 6000 chars of plain text
        return "\n".join(p.text for p in doc.paragraphs)[:max_chars]
    else:
        return file_bytes.decode("utf-8", errors="replace")[:max_chars]


def ingest_text(
    text: str,
    source_label: str,
    collection: str = LIBRARY_COLLECTION,
    chunk_type: ChunkType = "content",
    db: Optional[ChromaDBClient] = None,
    session_id: str = "",
    document_id: str = "",
) -> int:
    if db is None:
        db = ChromaDBClient()
    chunks = chunk_text(text)
    if not chunks:
        return 0
    ids = [f"{source_label}_{i}_{uuid.uuid4().hex[:6]}" for i in range(len(chunks))]
    metadatas = [
        {
            "source": source_label,
            "chunk_index": i,
            "type": chunk_type,
            "session_id": session_id,
            "document_id": document_id,
        }
        for i in range(len(chunks))
    ]
    db.upsert(collection, chunks, metadatas, ids)
    return len(chunks)


def ingest_file(
    file_bytes: bytes,
    filename: str,
    collection: str = LIBRARY_COLLECTION,
    chunk_type: ChunkType = "content",
    db: Optional[ChromaDBClient] = None,
    skip_if_indexed: bool = True,
    session_id: str = "",
) -> int:
    """Chunk and index a file. Returns chunk count (0 if already indexed and skip_if_indexed=True).

    Chunks are tagged with session_id + document_id (content hash) so retrieval can
    scope to exactly the current session's content and never bleed across PDFs.
    """
    if db is None:
        db = ChromaDBClient()

    content_hash = db.file_hash(file_bytes)
    if skip_if_indexed and db.is_indexed(content_hash):
        return 0

    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        text = _extract_pdf(file_bytes)
    elif ext == "docx":
        text = _extract_docx(file_bytes)
    else:
        text = file_bytes.decode("utf-8", errors="replace")

    count = ingest_text(
        text, filename, collection, chunk_type=chunk_type, db=db,
        session_id=session_id, document_id=content_hash,
    )
    if count > 0:
        db.mark_indexed(content_hash, filename)
    return count
